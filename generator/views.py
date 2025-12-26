from django.shortcuts import render
from django.http import HttpResponse

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ---------------- HELPERS ----------------

def fixed_height(row, cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(int(cm * 567)))
    h.set(qn("w:hRule"), "exact")
    trPr.append(h)


def set_columns(table):
    table.autofit = False
    widths = [Cm(1.2), Cm(4.8), Cm(11.0)]
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w
            row.cells[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP


def bold(paragraph):
    for r in paragraph.runs:
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)


# ---------------- VIEW ----------------

def generate_doc(request):

    if request.method == "POST":

        # ===== Applicant (HTML → Backend mapping) =====
        client_name = request.POST.get("client_name", "")
        branch_address = request.POST.get("branch_address", "")
        correspondence_address = request.POST.get("correspondence_address", "")
        telephone_no = request.POST.get("telephone_no", "")
        mobile = request.POST.get("mobile", "")
        email = request.POST.get("email", "")

        # ===== Opposite Party =====
        op_name = request.POST.get("customer_name", "")
        op_reg_address = request.POST.get("op_registered_address", "")
        op_corr_address = request.POST.get("op_correspondence_address", "")
        op_telephone = request.POST.get("op_telephone", "")
        op_mobile = request.POST.get("op_mobile", "")
        op_email = request.POST.get("op_email", "")

        # ===== Dispute =====
        dispute_nature = request.POST.get("dispute_nature", "")

        doc = Document()

        # -------- Page margins --------
        sec = doc.sections[0]
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2)
        sec.right_margin = Cm(2)

        # -------- HEADER --------
        p = doc.add_paragraph(
            "FORM ‘A’\n"
            "MEDIATION APPLICATION FORM\n"
            "[REFER RULE 3(1)]"
        )
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bold(p)

        p = doc.add_paragraph(
            "Mumbai District Legal Services Authority\n"
            "City Civil Court, Mumbai"
        )
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # -------- TABLE --------
        table = doc.add_table(rows=0, cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # DETAILS OF PARTIES
        row = table.add_row()
        fixed_height(row, 0.9)
        row.cells[0].merge(row.cells[2])
        row.cells[0].paragraphs[0].add_run("DETAILS OF PARTIES:").bold = True

        # 1. Name of Applicant
        row = table.add_row()
        fixed_height(row, 0.9)
        row.cells[0].text = "1"
        row.cells[1].text = "Name of Applicant"
        row.cells[2].text = client_name
        bold(row.cells[1].paragraphs[0])

        # Address and contact details of Applicant
        row = table.add_row()
        fixed_height(row, 0.9)
        row.cells[1].merge(row.cells[2])
        row.cells[1].paragraphs[0].add_run(
            "Address and contact details of Applicant"
        ).bold = True

        # Applicant Address
        row = table.add_row()
        fixed_height(row, 3.0)
        row.cells[1].text = "Address"
        bold(row.cells[1].paragraphs[0])

        cell = row.cells[2]
        p = cell.paragraphs[0]
        p.add_run("REGISTERED ADDRESS:\n").bold = True
        p.add_run(branch_address + "\n\n")
        p.add_run("CORRESPONDENCE BRANCH ADDRESS:\n").bold = True
        p.add_run(correspondence_address)

        # Telephone No. (Applicant)
        row = table.add_row()
        fixed_height(row, 0.8)
        row.cells[1].text = "Telephone No."
        row.cells[2].text = telephone_no
        bold(row.cells[1].paragraphs[0])

        # Mobile No. (Applicant)
        row = table.add_row()
        fixed_height(row, 0.8)
        row.cells[1].text = "Mobile No."
        row.cells[2].text = mobile
        bold(row.cells[1].paragraphs[0])

        # Email ID (Applicant)
        row = table.add_row()
        fixed_height(row, 0.8)
        row.cells[1].text = "Email ID"
        row.cells[2].text = email
        bold(row.cells[1].paragraphs[0])

        # 2. Opposite Party Header
        row = table.add_row()
        fixed_height(row, 0.9)
        row.cells[0].merge(row.cells[2])
        row.cells[0].paragraphs[0].add_run(
            "2  Name, Address and Contact details of Opposite Party:"
        ).bold = True

        # Name (Opposite Party)
        row = table.add_row()
        fixed_height(row, 0.9)
        row.cells[1].text = "Name"
        row.cells[2].text = op_name
        bold(row.cells[1].paragraphs[0])

        # Address (Opposite Party)
        row = table.add_row()
        fixed_height(row, 3.0)
        row.cells[1].text = "Address"
        bold(row.cells[1].paragraphs[0])

        cell = row.cells[2]
        p = cell.paragraphs[0]
        p.add_run("REGISTERED ADDRESS:\n").bold = True
        p.add_run(op_reg_address + "\n\n")
        p.add_run("CORRESPONDENCE ADDRESS:\n").bold = True
        p.add_run(op_corr_address)

        # Telephone (Opposite Party)
        row = table.add_row()
        fixed_height(row, 0.8)
        row.cells[1].text = "Telephone No."
        row.cells[2].text = op_telephone
        bold(row.cells[1].paragraphs[0])

        # Mobile (Opposite Party)
        row = table.add_row()
        fixed_height(row, 0.8)
        row.cells[1].text = "Mobile No."
        row.cells[2].text = op_mobile
        bold(row.cells[1].paragraphs[0])

        # Email (Opposite Party)
        row = table.add_row()
        fixed_height(row, 0.8)
        row.cells[1].text = "Email ID"
        row.cells[2].text = op_email
        bold(row.cells[1].paragraphs[0])

        # DETAILS OF DISPUTE
        row = table.add_row()
        fixed_height(row, 0.9)
        row.cells[0].merge(row.cells[2])
        row.cells[0].paragraphs[0].add_run("DETAILS OF DISPUTE:").bold = True

        # Centered rule
        row = table.add_row()
        fixed_height(row, 0.9)
        row.cells[0].text = ""
        merged = row.cells[1].merge(row.cells[2])
        p = merged.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(
            "THE COMMERCIAL COURTS (PRE-INSTITUTION SETTLEMENT) RULES, 2018"
        )
        r.bold = True
        r.underline = True

        # Nature of dispute (AUTO HEIGHT)
        row = table.add_row()
        row.cells[0].text = ""
        merged = row.cells[1].merge(row.cells[2])
        p = merged.paragraphs[0]
        p.add_run(
            "Nature of disputes as per section 2(1)(c) of the Commercial Courts Act, "
            "2015 (4 of 2016):\n"
        ).bold = True
        p.add_run(dispute_nature)

        set_columns(table)

        response = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response["Content-Disposition"] = (
            "attachment; filename=FORM_A_Mediation_Application.docx"
        )
        doc.save(response)
        return response

    return render(request, "form.html")
